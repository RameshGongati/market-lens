"""Build the Word research report from the aggregated backtest outputs."""
from __future__ import annotations

import glob
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
OUT = REPO_ROOT / "research_engine" / "output"
CHARTS = OUT / "charts"

TF_ORDER = ["daily", "weekly", "60m", "75m", "15m"]
PRETTY = {
    "gap_up_go": "Gap-Up Continuation", "gap_down_go": "Gap-Down Continuation",
    "zone_touch_fresh": "Fresh Zone Touch (Type-1 limit)", "demand_bounce": "Demand Zone Bounce",
    "supply_rejection": "Supply Zone Rejection", "ema20_bounce": "EMA20 Bounce",
    "ema20_rejection": "EMA20 Rejection", "sma50_pullback_long": "SMA50 Pullback Long",
    "sma50_pullback_short": "SMA50 Pullback Short", "macd_cross_long": "MACD Bull Cross",
    "macd_cross_short": "MACD Bear Cross", "bb_squeeze_breakout": "Bollinger Squeeze Breakout",
    "bb_squeeze_breakdown": "Bollinger Squeeze Breakdown", "rsi_bull_divergence": "RSI Bullish Divergence",
    "rsi_bear_divergence": "RSI Bearish Divergence", "stoch_bull_divergence": "Stochastic Bullish Divergence",
    "stoch_bear_divergence": "Stochastic Bearish Divergence", "hh_hl_continuation": "HH/HL Continuation",
    "lh_ll_continuation": "LH/LL Continuation", "pdh_breakout": "Prev-Day-High Breakout",
    "pdl_breakdown": "Prev-Day-Low Breakdown", "fib_pullback_long": "Fibonacci Pullback Long",
    "fib_pullback_short": "Fibonacci Pullback Short", "bullish_engulfing": "Bullish Engulfing",
    "bearish_engulfing": "Bearish Engulfing", "hammer": "Hammer", "shooting_star": "Shooting Star",
    "inverted_hammer": "Inverted Hammer", "morning_star": "Morning Star", "evening_star": "Evening Star",
    "inside_bar_breakout": "Inside Bar Breakout", "inside_bar_breakdown": "Inside Bar Breakdown",
    "nr7_breakout": "NR7 Breakout", "nr7_breakdown": "NR7 Breakdown",
    "strong_bull_candle": "Strong Bullish Candle (M5)", "strong_bear_candle": "Strong Bearish Candle (M5)",
    "triangle_sym_breakout": "Symmetrical Triangle Breakout", "triangle_sym_breakdown": "Symmetrical Triangle Breakdown",
    "ascending_triangle_break": "Ascending Triangle Breakout", "descending_triangle_break": "Descending Triangle Breakdown",
    "vcp_breakout": "VCP / Tight Base Breakout", "range_breakout": "Rectangle Range Breakout",
    "range_breakdown": "Rectangle Range Breakdown", "bull_flag_break": "Bull Flag Breakout",
    "bear_flag_break": "Bear Flag Breakdown", "bull_pennant_break": "Bull Pennant Breakout",
    "bear_pennant_break": "Bear Pennant Breakdown", "double_bottom_break": "Double Bottom Breakout",
    "double_top_break": "Double Top Breakdown",
}


def pretty(s: str) -> str:
    return PRETTY.get(s, s)


# ----------------------------------------------------------------------- doc helpers
def add_heading(doc, text, level):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_table(doc, header, rows, widths=None, font_size=8):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for j, htxt in enumerate(header):
        hdr[j].text = ""
        r = hdr[j].paragraphs[0].add_run(str(htxt))
        r.bold = True
        r.font.size = Pt(font_size)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run("" if v is None or (isinstance(v, float) and np.isnan(v)) else str(v))
            r.font.size = Pt(font_size)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    return t


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    inner = OxmlElement("w:t")
    inner.text = "Table of Contents — in Word: right-click here and choose 'Update Field'."
    r = OxmlElement("w:r")
    r.append(inner)
    fld.append(r)
    run._r.addnext(fld)


def fmt_row(rk, cols):
    out = []
    for c in cols:
        v = rk.get(c)
        if isinstance(v, float):
            if c in ("n", "n_symbols", "symbols_with_3plus", "n_with", "n_without"):
                v = int(v)
            else:
                v = round(v, 3) if abs(v) < 10 else round(v, 1)
        out.append(v)
    return out


RANK_COLS = ["setup_name", "timeframe", "n", "expectancy_r", "win_rate_2r", "hit_1r_pct",
             "profit_factor", "avg_mfe_r", "avg_mae_r", "false_breakout_pct",
             "avg_holding_bars", "pct_symbols_positive"]
RANK_HDRS = ["Setup", "TF", "N", "Exp (R)", "2R win %", "1R hit %", "PF", "MFE (R)",
             "MAE (R)", "False brk %", "Hold (bars)", "% syms +ve"]


def ranking_table(doc, r, mask, top=None, sort="expectancy_r", asc=False):
    sub = r[mask].sort_values(sort, ascending=asc)
    if top:
        sub = sub.head(top)
    rows = []
    for _, rk in sub.iterrows():
        d = rk.to_dict()
        d["setup_name"] = pretty(d["setup_name"]) + (" (short)" if d.get("bullish_or_bearish") == "bearish" else "")
        rows.append(fmt_row(d, RANK_COLS))
    add_table(doc, RANK_HDRS, rows,
              widths=[1.9, 0.5, 0.55, 0.6, 0.6, 0.6, 0.5, 0.55, 0.55, 0.6, 0.6, 0.6])


def main() -> None:
    r = pd.read_csv(OUT / "setup_rankings.csv")
    combos = pd.read_csv(OUT / "combo_rankings.csv")
    flags = pd.read_csv(OUT / "flag_effects.csv")
    sectors = pd.read_csv(OUT / "sector_findings.csv")
    stocks = pd.read_csv(OUT / "stock_findings.csv")
    trades = pd.concat([pd.read_parquet(p) for p in sorted(glob.glob(str(OUT / "trades_*_all.parquet")))],
                       ignore_index=True)
    n_total = len(trades)

    doc = Document()
    for sec in doc.sections:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(0.7)

    # ---------------- title
    t = add_para(doc, "Indian F&O Chart Pattern & Indicator Research", bold=True, size=24)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = add_para(doc, "One-year historical backtest across the NSE F&O universe — evidence base for Market Lens scanner rules",
                 italic=True, size=13)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "")
    meta = add_para(doc, f"Prepared 13 Aug 2026  ·  {n_total:,} simulated trades  ·  208 F&O stocks + NIFTY 50 + BANK NIFTY  ·  5 timeframes", size=10)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warn = add_para(doc, "For education and research only. Historical performance does not guarantee future results. "
                         "Nothing in this document is investment advice or a recommendation to trade.", italic=True, size=9)
    warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_heading(doc, "Contents", 1)
    add_toc(doc)
    doc.add_page_break()

    # ---------------- 1. executive summary
    add_heading(doc, "1. Executive Summary", 1)
    add_para(doc, "This study simulated every occurrence of 46 setups (candlestick, chart pattern, indicator and "
                  "demand/supply-zone based, long and short) across 210 instruments and five timeframes over the last "
                  "year, under one uniform, conservative trade model: enter on the next bar, stop at a setup-specific "
                  "level with an ATR buffer, take profit at 2R, time-stop otherwise, 0.1% round-trip costs, and any bar "
                  "that touches both stop and target counted as a loss.")
    add_bullets(doc, [
        "Most setups do NOT survive costs under a mechanical 2R rule. Out of 46 setups, only a handful show positive "
        "expectancy at reliable sample sizes. This is the single most important finding: pattern names alone carry "
        "almost no edge — context and selection do.",
        "Best bullish setup: Gap-Up Continuation (open ≥1.3% above prior high, close holds) — +0.37R expectancy, "
        "profit factor 2.16, 1,553 trades, positive on 77% of symbols, and positive on every timeframe tested.",
        "Best zone setup: Fresh Zone Touch (GTF Type-1 limit at the proximal of a never-tested zone) — +0.12R over "
        "5,272 long-side trades (34% reach 2R; a 2R target only needs ~33% to break even). The demand-side "
        "confirmation bounce (+0.04R) also clears zero; the supply-side equivalents are weaker, consistent with a "
        "rising market year.",
        "Higher timeframes beat lower ones decisively: weekly bullish setups averaged +0.01R, daily −0.07R, and 15m "
        "−0.34R. Intraday noise plus costs destroy most edges below the hourly timeframe.",
        "Bearish setups broadly failed (avg −0.05R to −0.27R by timeframe): the test year trended up. Shorts need "
        "stronger filters, and several only work near supply zones.",
        "Context flags change outcomes more than setup choice: EMA20 Rejection shorts on 75m go from −0.16R to "
        "+0.20R when taken at a live supply zone; hammers on 15m improve by +0.41R near demand (though still "
        "negative); symmetrical-triangle breakouts on 60m flip positive with Fibonacci confluence.",
        "A counterintuitive regime result: bullish signals taken while NIFTY was already above its 20-EMA performed "
        "WORSE (−0.25R) than those taken below it (−0.10R) — chasing an extended market was costlier than buying "
        "fear. A simple 'market must be bullish' filter would have hurt, not helped.",
    ])
    add_para(doc, "Recommended for Market Lens (detail in §15): a Gap-from-Zone scanner, SMA50-trend gating on the "
                  "existing fresh-zone logic, a demand-side-only default for the confirmation screener, a weekly "
                  "bullish-reversal scanner, and a 75m supply-rejection short scanner. Recommended against: standalone "
                  "Bollinger-squeeze, flag/pennant and inverted-hammer rules, and any of these setups on raw index "
                  "underlyings.")

    # ---------------- 2. methodology
    doc.add_page_break()
    add_heading(doc, "2. Methodology", 1)
    add_heading(doc, "2.1 Universe and period", 2)
    add_bullets(doc, [
        "208 NSE F&O stocks (Market Lens F&O watchlist) plus NIFTY 50 and BANK NIFTY as tradeable instruments.",
        "Sector context from 10 NSE sector indices; market context from NIFTY 50 vs its 20-EMA.",
        "Test window: 11 Aug 2025 – 11 Aug 2026 for daily/weekly (2 years fetched for indicator warm-up); "
        "the full available year for 60m; ~60 trading days for 75m and 15m (provider limit, stated in §3).",
        "149 of 208 stocks carry a sector label from the shipped sector lists; the rest are grouped as 'Other'.",
    ])
    add_heading(doc, "2.2 Trade model (identical for every setup)", 2)
    add_table(doc, ["Rule", "Value"], [
        ["Entry", "Open of the bar AFTER the signal bar closes. Exception: fresh-zone touches fill at the zone's proximal (limit order), with no favourable credit on the fill bar."],
        ["Stop-loss", "Setup-specific structure level (pattern boundary, candle extreme, zone distal) ± 0.1×ATR(14) buffer; breakout stops 1.0×ATR beyond the broken level."],
        ["Target", "2R primary (evaluated first-touch). 1R and 3R hit rates tracked from excursion for context."],
        ["Time stop", "20 bars daily · 8 weekly · 35×60m · 30×75m · 75×15m — exit at close."],
        ["Same-bar conflict", "Bar touches stop AND target → counted as LOSS (intrabar order unknowable)."],
        ["Costs", "0.1% of entry per round trip (brokerage + slippage)."],
        ["Degenerate stops", "Risk <0.1% or >12% of entry → trade skipped."],
        ["Duplicates", "Same setup re-firing within 3 bars suppressed; same pattern-breakout within 10 bars suppressed."],
    ], widths=[1.4, 5.7], font_size=9)
    add_heading(doc, "2.3 Look-ahead controls", 2)
    add_bullets(doc, [
        "All indicators computed point-in-time; swing points only count once confirmed 3 bars later.",
        "Zone detection uses the Market Lens engine, but with its forward-looking invalidation filter disabled for "
        "research: the engine normally discards zones that get broken LATER, which would leave only zones whose "
        "stops were never hit (survivorship bias). An early run of this study showed exactly that failure — zero "
        "losing zone trades — and was rebuilt. Touches, invalidation and habitation are re-derived bar by bar.",
        "The still-forming bar is dropped everywhere, matching Market Lens behaviour.",
        "Market/sector regime for intraday signals uses the PREVIOUS day's index close (the same-day daily bar is "
        "not complete during the session).",
    ])
    add_heading(doc, "2.4 Definition of 'worked' and sample-size tiers", 2)
    add_para(doc, "A setup 'worked' only if it produced positive expectancy (mean R per trade, after costs) under the "
                  "model above. Win rate alone is not used: a 2R target converts to breakeven at ~33% wins. Tiers: "
                  "n ≥ 30 = reliable; 10–29 = indicative; < 10 = anecdote (never ranked). Consistency = share of "
                  "symbols (≥3 trades) with positive expectancy.")

    # ---------------- 3. data sources & limitations
    add_heading(doc, "3. Data Sources, Assumptions and Limitations", 1)
    add_bullets(doc, [
        "Prices: Yahoo Finance, UNADJUSTED (matching Market Lens zone levels). Yahoo occasionally omits whole "
        "sessions for some symbols; a missing bar can shift zone/base counts.",
        "Intraday history is capped by the provider: 60 days for 15m (and therefore 75m, which is resampled from "
        "15m). Their results cover ~2.5 months, not a full year — treat as indicative.",
        "125-minute bars were NOT tested: no native interval exists and the only resample source (5m) is capped at "
        "60 days — too small a sample to rank fairly.",
        "Index instruments carry no volume on Yahoo; volume-based flags are always false for NIFTY/BANKNIFTY rows.",
        "Options were NOT simulated. Signals are on the underlying; option pricing (IV, theta, strikes) is out of "
        "scope. 'Options-style' here means directional setups with defined risk suitable for option entries.",
        "Weekly = 52 bars/instrument → small per-symbol samples; weekly rows lean on cross-sectional breadth.",
        "One year, one market regime (broadly upward). Bearish-setup readings in particular may not generalise.",
        "No portfolio effects: trades simulated independently; no capital constraints, margin or position sizing.",
    ])

    # ---------------- 4. overall results
    doc.add_page_break()
    add_heading(doc, "4. Overall Results", 1)
    per_tf = trades.groupby(["timeframe", "bullish_or_bearish"])["r_multiple"].agg(["size", "mean"]).round(3)
    rows = []
    for tf in TF_ORDER:
        for side in ("bullish", "bearish"):
            try:
                size, mean = per_tf.loc[(tf, side)]
                rows.append([tf, side, int(size), round(mean, 3)])
            except KeyError:
                pass
    add_heading(doc, "4.1 Trades and average expectancy by timeframe and direction", 2)
    add_table(doc, ["Timeframe", "Direction", "Trades", "Avg R (all setups)"], rows, font_size=9)
    add_para(doc, "Averages across ALL setups are negative — most raw signals lose after costs. Rankings below "
                  "identify the exceptions. Daily and weekly clearly dominate intraday: the 15m columns absorb the "
                  "most noise and the most cost per unit of edge.", italic=True)
    add_heading(doc, "4.2 Best bullish setups (all timeframes pooled, n ≥ 100)", 2)
    ranking_table(doc, r, (r.timeframe == "ALL") & (r.n >= 100) & (r.bullish_or_bearish == "bullish"), top=10)
    add_heading(doc, "4.3 Best bearish setups (all timeframes pooled, n ≥ 100)", 2)
    ranking_table(doc, r, (r.timeframe == "ALL") & (r.n >= 100) & (r.bullish_or_bearish == "bearish"), top=10)
    add_para(doc, "Note how much weaker the best shorts are than the best longs this year. Only fresh supply-zone "
                  "touches and gap-down continuation approach breakeven at scale; classic bearish patterns "
                  "(flags, LH/LL continuation) rank at the bottom of the whole study.", italic=True)
    add_heading(doc, "4.4 Best setups per timeframe (n ≥ 50)", 2)
    for tf in TF_ORDER:
        add_para(doc, tf.upper(), bold=True)
        ranking_table(doc, r, (r.timeframe == tf) & (r.n >= 50), top=5)
    add_heading(doc, "4.5 Market regime effect", 2)
    reg = trades[trades.bullish_or_bearish == "bullish"].groupby("market_confirmation")["r_multiple"].agg(["size", "mean"]).round(3)
    rows = [[("NIFTY above 20-EMA" if k else "NIFTY below 20-EMA"), int(v["size"]), v["mean"]] for k, v in reg.iterrows()]
    add_table(doc, ["Regime at signal", "Bullish trades", "Avg R"], rows, font_size=9)
    add_para(doc, "Bullish signals in an already-extended market underperformed. The practical reading: regime "
                  "filters should gate AGGRESSION (position size, target ambition), not signal existence — and "
                  "'buy strength' worked less well than 'buy fear' for these mechanical entries this year.", italic=True)

    # ---------------- 5-8. per-category analysis
    doc.add_page_break()
    add_heading(doc, "5. Demand/Supply Zone Setups (GTF engine)", 1)
    ranking_table(doc, r, (r.timeframe == "ALL") & (r.setup_name.isin(["zone_touch_fresh", "demand_bounce", "supply_rejection"])))
    add_para(doc, "These use Market Lens's own zone engine, so the numbers transfer directly to the app. The fresh-zone "
                  "limit touch is the strongest programmatic edge in the study at scale (+0.12R long / +0.06R short, "
                  ">10,000 trades, ~66% of symbols positive). The confirmation bounce (the app's Zone Confirmation "
                  "Screener concept) is mildly positive on the demand side and negative on the supply side this year. "
                  "Per-timeframe: both concentrate their edge on daily and 60m; 15m dilutes it.")
    add_heading(doc, "6. Chart Pattern Setups", 1)
    pat_names = [k for k, v in PRETTY.items() if "Triangle" in v or "Flag" in v or "Pennant" in v
                 or "VCP" in v or "Rectangle" in v or "Double" in v]
    ranking_table(doc, r, (r.timeframe == "ALL") & (r.setup_name.isin(pat_names)) & (r.n >= 50))
    add_para(doc, "Raw pattern breakouts hover at or below zero: false-breakout rates of 14–24% (closing back inside "
                  "within 3 bars without reaching 1R) eat the edge. Two exceptions worth noting: bull pennants are the "
                  "only pattern family positive at ALL-timeframe level, and symmetrical-triangle breakouts on 60m turn "
                  "positive when a Fibonacci level sits in the pattern (§9). Wedges, channels, cup-and-handle and "
                  "head-and-shoulders were NOT tested — Market Lens has no detector for them and building validated "
                  "ones was out of scope; treat their absence as untested, not as failed.")
    add_heading(doc, "7. Candlestick Setups", 1)
    candle_names = ["bullish_engulfing", "bearish_engulfing", "hammer", "shooting_star", "inverted_hammer",
                    "morning_star", "evening_star", "inside_bar_breakout", "inside_bar_breakdown",
                    "nr7_breakout", "nr7_breakdown", "strong_bull_candle", "strong_bear_candle"]
    ranking_table(doc, r, (r.timeframe == "ALL") & (r.setup_name.isin(candle_names)))
    add_para(doc, "Standalone candlestick patterns are net losers at every reliable sample size — the inverted hammer "
                  "is the single worst setup in the study (−0.52R over 31,655 trades). They only become useful as "
                  "CONFIRMATION at a location (see §9): the same hammer gains +0.4R when it prints inside a live "
                  "demand zone. Morning/evening stars and hammers on WEEKLY bars are the exception: weekly reversal "
                  "candles were mildly positive (morning star +0.17R, hammer +0.11R).")
    add_heading(doc, "8. Indicator Setups", 1)
    ind_names = ["ema20_bounce", "ema20_rejection", "sma50_pullback_long", "sma50_pullback_short",
                 "macd_cross_long", "macd_cross_short", "bb_squeeze_breakout", "bb_squeeze_breakdown",
                 "rsi_bull_divergence", "rsi_bear_divergence", "stoch_bull_divergence", "stoch_bear_divergence",
                 "gap_up_go", "gap_down_go", "hh_hl_continuation", "lh_ll_continuation",
                 "pdh_breakout", "pdl_breakdown", "fib_pullback_long", "fib_pullback_short"]
    ranking_table(doc, r, (r.timeframe == "ALL") & (r.setup_name.isin(ind_names)) & (r.n >= 100))
    add_para(doc, "Gap-up continuation is the standout of the entire study — positive on every timeframe (daily "
                  "+0.41R, 75m +0.40R, 15m +0.37R, 60m +0.34R), profit factor >2 everywhere, 77% of symbols "
                  "positive. Its bearish mirror is roughly breakeven. RSI/stochastic bullish divergences are small "
                  "positives on daily; MACD bull crosses work on weekly (+0.28R) and 75m (+0.14R) but not daily. "
                  "Bollinger-squeeze breakouts, fib pullbacks taken blind, SMA50 pullbacks and PDH/PDL breaks all "
                  "lose after costs as standalone rules.")

    # ---------------- 9. combinations
    doc.add_page_break()
    add_heading(doc, "9. Combination Setups", 1)
    add_para(doc, "Each named combination filters a base setup by context flags recorded at signal time. "
                  "'Δ vs base' is the expectancy change against the unfiltered setup — the honest measure of whether "
                  "the extra condition earns its keep.")
    ct = combos[(combos.timeframe == "ALL") & (combos.n >= 30)].sort_values("expectancy_r", ascending=False)
    rows = [[c["combo"], int(c["n"]), c["expectancy_r"], c["base_expectancy_r"], c["delta_expectancy_r"],
             c["win_rate_2r"], c["profit_factor"]] for _, c in ct.iterrows()]
    add_table(doc, ["Combination", "N", "Exp (R)", "Base exp", "Δ vs base", "2R win %", "PF"], rows,
              widths=[2.5, 0.5, 0.6, 0.6, 0.6, 0.7, 0.5], font_size=8)
    add_bullets(doc, [
        "Helpful: SMA50 trend alignment on fresh-zone touches (+0.04Δ on an already-positive base — the best "
        "combined rule at scale: +0.13R, n=4,439); Fibonacci confluence on demand bounces (+0.01Δ); volume "
        "expansion on double-top breaks and prior contraction on 15m range breakdowns (§10).",
        "Harmful or neutral, despite sounding right: requiring NIFTY 'market support' on demand bounces "
        "(−0.08Δ) or on range breakouts (−0.06Δ); volume expansion on demand bounces (−0.06Δ); EMA20 confluence "
        "on demand bounces (−0.05Δ). Filters must be tested, not assumed.",
        "Zone location rescues weak setups without always making them tradeable: bearish engulfing near supply "
        "improves +0.10Δ but stays negative; the 75m EMA20-rejection-at-supply is the one weak-base combo that "
        "turns clearly positive (+0.20R, n=105 — indicative).",
    ])
    add_heading(doc, "9.1 Strongest single-flag effects (n ≥ 100 both sides)", 2)
    f2 = flags[(flags.n_with >= 100) & (flags.n_without >= 100)].sort_values("delta", ascending=False).head(10)
    rows = [[pretty(x["setup_name"]), x["timeframe"], x["flag"], int(x["n_with"]), x["exp_with"],
             x["exp_without"], x["delta"]] for _, x in f2.iterrows()]
    add_table(doc, ["Setup", "TF", "Flag", "N (with)", "Exp with", "Exp without", "Δ"], rows,
              widths=[1.7, 0.5, 1.5, 0.7, 0.7, 0.8, 0.6], font_size=8)

    # ---------------- 10-11. stocks & sectors
    doc.add_page_break()
    add_heading(doc, "10. Stock-Specific Findings", 1)
    add_para(doc, "Per-stock expectancy across all setups separates instruments where mechanical setups behaved from "
                  "those that chopped. High-beta names with clean trends (SWIGGY, KAYNES, PAYTM, IDEA, NBCC, "
                  "ANGELONE) top the table; defensive/rangebound names (MARICO, LUPIN, ALKEM, POWERGRID) and the "
                  "INDEX UNDERLYINGS THEMSELVES sit at the bottom — BANKNIFTY was the second-worst instrument in the "
                  "study. For options-style scanning, prefer the volatile trending names and treat index signals as "
                  "context, not trades.")
    top10, bot10 = stocks.head(10), stocks.tail(10)
    for label, chunk in (("Most setup-friendly stocks", top10), ("Least setup-friendly stocks", bot10)):
        add_para(doc, label, bold=True)
        rows = [[s["symbol"], s["sector"], int(s["n_trades"]), s["best_setup"],
                 s["bullish_exp_r"], s["bearish_exp_r"], s["overall_exp_r"]] for _, s in chunk.iterrows()]
        add_table(doc, ["Symbol", "Sector", "Trades", "Best setup (n≥5)", "Bull exp", "Bear exp", "Overall"],
                  rows, widths=[0.9, 1.1, 0.6, 2.4, 0.7, 0.7, 0.7], font_size=8)
    add_para(doc, "The full per-stock table (210 rows) is in stock_findings.csv.", italic=True)

    add_heading(doc, "11. Sector-Specific Findings", 1)
    rows = [[s["sector"], int(s["n_trades"]), int(s["n_symbols"]), pretty(str(s["best_setup"])),
             s["best_setup_exp_r"], s["best_timeframe"], s["bullish_exp_r"], s["bearish_exp_r"]]
            for _, s in sectors.iterrows()]
    add_table(doc, ["Sector", "Trades", "Syms", "Best setup (n≥20)", "Best exp", "Best TF", "Bull exp", "Bear exp"],
              rows, widths=[1.3, 0.7, 0.5, 1.9, 0.7, 0.7, 0.7, 0.7], font_size=8)
    add_para(doc, "Realty, Consumer Durables and IT tolerated mechanical setups best; Pharma and FMCG were the most "
                  "hostile (choppy, mean-reverting). Gap-up continuation is the best setup in 8 of 14 groups — its "
                  "edge is broad, not sector-specific. Weekly is the best timeframe in most sectors, reinforcing §4.",
             italic=True)

    # ---------------- 12. worked examples
    doc.add_page_break()
    add_heading(doc, "12. Worked and Failed Examples (charts)", 1)
    add_para(doc, "One winning and one losing example per headline setup — failures are as instructive as wins. "
                  "Green line = entry, red dashed = stop, purple dotted = 2R target; grey vertical = signal bar; "
                  "blue = EMA20, orange = SMA50.")
    charts = sorted(CHARTS.glob("*.png"))
    tmap = {}
    for _, tr in trades.iterrows():
        key = (tr["symbol"], tr["timeframe"], tr["setup_name"], str(tr["signal_date"])[:10], tr["result"])
        tmap.setdefault(key, tr)
    for png in charts:
        parts = png.stem.split("_")
        tag = parts[0]
        date = parts[-1]
        sym = parts[1]
        tf = parts[2]
        setup = "_".join(parts[3:-1])
        row = None
        for res in (("win", "timeout") if tag == "worked" else ("loss",)):
            row = tmap.get((sym, tf, setup, date, res))
            if row is not None:
                break
        doc.add_picture(str(png), width=Inches(6.9))
        if row is not None:
            why = ("Price reached 2R before the stop — favourable structure held." if tag == "worked"
                   else "Stop hit before 2R — the move failed after entry.")
            extra = ""
            if row.get("near_demand"):
                extra += " Signal printed near a live demand zone."
            if row.get("near_supply"):
                extra += " Signal printed near a live supply zone."
            if row.get("volume_confirmation"):
                extra += " Volume expanded on the signal bar."
            if row.get("market_confirmation") is False and row["bullish_or_bearish"] == "bullish":
                extra += " Taken while NIFTY was below its 20-EMA."
            cap = (f"{sym} · {tf} · {pretty(setup)} ({row['bullish_or_bearish']}) · signal {date} · "
                   f"entry {row['entry_price']} / stop {row['stop_loss']} / 2R target {row['target_2']} · "
                   f"outcome: {row['result'].upper()} {row['r_multiple']:+.2f}R in {int(row['holding_period'])} bars. {why}{extra}")
        else:
            cap = f"{sym} · {tf} · {pretty(setup)} · signal {date} · {'reached 2R' if tag=='worked' else 'stopped out'}."
        add_para(doc, cap, size=9, italic=True)

    # ---------------- 13. false signals & traps
    doc.add_page_break()
    add_heading(doc, "13. False Signals and Traps", 1)
    add_heading(doc, "13.1 Worst setups (n ≥ 1,000 — avoid as standalone rules)", 2)
    ranking_table(doc, r, (r.timeframe == "ALL") & (r.n >= 1000), top=8, sort="expectancy_r", asc=True)
    losers = trades[trades["result"] == "loss"]
    winners = trades[trades["hit_2r"]]
    ctx = []
    for label, flag in [("Volume expansion on signal", "volume_confirmation"),
                        ("Aligned with SMA200 trend", "sma200_trend"),
                        ("NIFTY regime aligned", "market_confirmation"),
                        ("Near opposing zone (long under supply / short over demand)", None)]:
        if flag == "sma200_trend":
            lv = (np.where(losers.bullish_or_bearish == "bullish", losers[flag] == "above", losers[flag] == "below")).mean()
            wv = (np.where(winners.bullish_or_bearish == "bullish", winners[flag] == "above", winners[flag] == "below")).mean()
        elif flag is None:
            lv = (np.where(losers.bullish_or_bearish == "bullish", losers["near_supply"], losers["near_demand"])).mean()
            wv = (np.where(winners.bullish_or_bearish == "bullish", winners["near_supply"], winners["near_demand"])).mean()
        else:
            lv = losers[flag].fillna(False).astype(bool).mean()
            wv = winners[flag].fillna(False).astype(bool).mean()
        ctx.append([label, f"{lv*100:.1f}%", f"{wv*100:.1f}%"])
    add_heading(doc, "13.2 Context of losers vs 2R winners (all trades)", 2)
    add_table(doc, ["Condition", "Losing trades", "2R-winning trades"], ctx, font_size=9)
    add_bullets(doc, [
        "Breakout traps: rectangle and pennant breaks closed back inside within 3 bars on ~24% of signals; double "
        "tops/bottoms ~14%. Requiring the FIRST close beyond the level to also be an M5-exciting candle, or waiting "
        "for a held retest, is the structural fix (Market Lens's held_retest already encodes this).",
        "Counter-trend knife-catching: hammer / inverted hammer after 3 down bars is the worst family — the decline "
        "usually continues. Only location (a live demand zone) made them usable, and even then mostly on 15m as "
        "scalps.",
        "Chasing extension: bullish entries with NIFTY above its 20-EMA underperformed by 0.15R — most mechanical "
        "signals late in a leg are exhaustion, not continuation.",
        "Squeeze fakeouts: Bollinger-squeeze breaks fail unless a zone sits behind the move (Δ +0.24R with zone "
        "alignment) — volatility compression alone says nothing about direction.",
        "Index underlyings: every setup class scored worse on NIFTY/BANKNIFTY than on stocks — mechanical "
        "pattern trades on the indices themselves were the weakest slice of the study.",
    ])

    # ---------------- 14. timeframes  (short, folded into 4; kept as anchor)
    add_heading(doc, "14. Timeframe Guidance", 1)
    add_bullets(doc, [
        "WEEKLY: best expectancy per signal; few signals (52 bars/yr). Use for swing bias and reversal candles.",
        "DAILY: the workhorse — zone setups, gap continuation and divergences all clear costs here.",
        "60m: acceptable for zone touches and gap continuation; pattern breakouts marginal.",
        "75m: interesting niche — MACD crosses and EMA20-rejection-at-supply positive; only ~60 days of data (indicative).",
        "15m: avoid for these mechanical rules — only gap continuation and zone-located hammers survived, thinly.",
    ])

    # ---------------- 15. scanner recommendations
    doc.add_page_break()
    add_heading(doc, "15. Market Lens Scanner Recommendations", 1)
    add_para(doc, "Ordered by evidence strength. 'Confidence logic' maps to the study's expectancy/consistency so the "
                  "UI can rank candidates the way the evidence does.")
    recs = [
        ("R1 — Gap-Up Continuation (NEW)", [
            ("Basis", "+0.37R, PF 2.16, n=1,553, 77% of symbols positive; works daily/60m/75m/15m"),
            ("Signal", "Open ≥ 1.3% above prior bar high (uses the existing M-gap threshold) AND close ≥ open"),
            ("Confirmation", "None required — this is the rare setup that works unfiltered. Optional: gap FROM a demand zone (tag exists)"),
            ("Avoid", "Bearish mirror is breakeven — ship long side only; skip if risk (prior-bar low distance) > 6%"),
            ("Stop / Target", "Prior bar low − 0.1×ATR / 2R primary, trail beyond"),
            ("Confidence", "Base 70; +10 if gap from live demand zone; +5 volume expansion; −10 if NIFTY >2% above 20-EMA"),
            ("Alert / Columns", "At open when gap detected · gap %, risk %, zone context, R to next supply"),
        ]),
        ("R2 — Fresh Zone Touch, trend-gated (UPGRADE to existing zones)", [
            ("Basis", "+0.12R n=5,272 long; SMA50 alignment lifts to +0.13R (best at-scale combo); 66% of symbols positive"),
            ("Signal", "Price touches proximal of a never-tested zone (existing engine logic)"),
            ("Confirmation", "SMA50 trend aligned (close > SMA50 for demand) — add as scanner gate or confidence boost"),
            ("Avoid", "Index underlyings; 15m; zones with risk >6%; do NOT add a NIFTY-bullish filter (tested negative)"),
            ("Stop / Target", "Distal − 0.1×ATR / 2R; opposing zone as stretch target"),
            ("Confidence", "Existing ODD score; +1 tier if SMA50 aligned (validated), not market regime"),
            ("Alert / Columns", "Existing proximity alert unchanged · add 'SMA50 aligned' column"),
        ]),
        ("R3 — Zone Confirmation Screener: demand-side default (TUNE existing)", [
            ("Basis", "Demand bounce +0.04R vs supply rejection −0.08R this year"),
            ("Signal", "Existing confirmation-close logic (enter+close-back-out)"),
            ("Confirmation", "Fibonacci level inside zone adds +0.01Δ (existing enhancer data)"),
            ("Avoid", "Default the screener to demand-only; keep supply behind a toggle labelled as counter-trend"),
            ("Stop / Target", "Unchanged (distal-buffered)"),
            ("Confidence", "Keep 4.5 grading; evidence here supports demand-side weighting"),
            ("Alert / Columns", "Unchanged"),
        ]),
        ("R4 — Weekly Bullish Reversal (NEW, low frequency)", [
            ("Basis", "Weekly hammer +0.11R n=411, morning star +0.17R n=190, MACD bull cross +0.28R n=145"),
            ("Signal", "On weekly bars: hammer or morning star, or MACD cross up below zero"),
            ("Confirmation", "Sector index above its 20-EMA (+0.24Δ on weekly strong-bull candles)"),
            ("Avoid", "Bearish weekly mirrors (all negative); stocks in Pharma/FMCG scored worst"),
            ("Stop / Target", "Pattern low − 0.1×ATR / 2R, 8-week time stop"),
            ("Confidence", "60 base; +10 sector aligned; +10 near weekly demand zone"),
            ("Alert / Columns", "Weekend scan after Friday close · pattern, sector regime, zone context"),
        ]),
        ("R5 — 75m Supply-Zone Rejection Short (NEW, indicative)", [
            ("Basis", "EMA20 rejection AT supply on 75m: +0.20R, n=105 (indicative tier; re-validate on more data)"),
            ("Signal", "75m bearish close rejecting EMA20 from below while inside/near a live supply zone"),
            ("Confirmation", "Price below SMA50 on 75m"),
            ("Avoid", "Without the zone the same signal is −0.16R — zone presence is mandatory, not optional"),
            ("Stop / Target", "max(bar high, EMA20) + 0.1×ATR / 2R"),
            ("Confidence", "50 base; +15 fresh zone; +10 SMA50 aligned; flag 'indicative sample' in UI"),
            ("Alert / Columns", "Intraday · zone freshness, distance to zone, 75m trend"),
        ]),
    ]
    for title, fields in recs:
        add_para(doc, title, bold=True)
        add_table(doc, ["Field", "Specification"], list(fields), widths=[1.2, 5.9], font_size=9)
    add_para(doc, "Do NOT build (evidence-based): standalone Bollinger-squeeze scanner (−0.06R/−0.13R), "
                  "flag/pennant breakout scanner (bear side −0.32R), inverted-hammer alerts (−0.52R), LH/LL "
                  "continuation shorts (−0.29R), fib-pullback entries without zone context (−0.31R), and any of "
                  "these rules applied to NIFTY/BANKNIFTY underlyings. These aren't 'not yet tested' — they were "
                  "tested and lost.", bold=True)

    add_heading(doc, "15.1 Final checklist", 2)
    add_table(doc, ["#", "Rule", "Action in Market Lens", "Evidence tier"], [
        ["R1", "Gap-Up Continuation", "New scanner rule + open-time alert", "Reliable (n=1,553)"],
        ["R2", "Fresh Zone Touch + SMA50 gate", "Add trend gate/boost to existing zone scan", "Reliable (n=4,439)"],
        ["R3", "Confirmation screener demand-default", "Config change + UI toggle", "Reliable (n=5,349)"],
        ["R4", "Weekly Bullish Reversal", "New weekly scan (weekend run)", "Reliable (n=746 pooled)"],
        ["R5", "75m Supply Rejection Short", "New intraday rule, flagged indicative", "Indicative (n=105)"],
        ["—", "Avoid-list above", "Do not implement", "Tested negative"],
    ], widths=[0.4, 2.2, 2.9, 1.6], font_size=9)

    # ---------------- 16. appendix
    doc.add_page_break()
    add_heading(doc, "16. Appendix — Full Ranking (per timeframe, n ≥ 50)", 1)
    for tf in TF_ORDER:
        add_para(doc, tf.upper(), bold=True)
        ranking_table(doc, r, (r.timeframe == tf) & (r.n >= 50))
    add_heading(doc, "16.1 Deliverables", 2)
    add_bullets(doc, [
        "signals_detailed.csv(.gz) — every simulated trade with entries/stops/targets/outcomes and context flags (287 MB raw).",
        "setup_rankings.csv — this report's ranking tables, all timeframes.",
        "combo_rankings.csv / flag_effects.csv — combination analysis.",
        "stock_findings.csv / sector_findings.csv — per-stock and per-sector tables.",
        "charts/ — 24 annotated worked/failed examples.",
        "research_engine/harness/ — the full reproducible harness (fetch → detect → simulate → aggregate).",
    ])

    dest = OUT / "FnO_Pattern_Research_Report.docx"
    doc.save(dest)
    print("saved", dest)


if __name__ == "__main__":
    main()
