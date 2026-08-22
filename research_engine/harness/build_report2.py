"""Word report #2 — Market Lens Overall Strategy Engine research."""
from __future__ import annotations

import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

REPO_ROOT = Path(__file__).resolve().parents[2]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from research_engine.harness.build_report import (  # noqa: E402
    add_bullets, add_heading, add_para, add_table, add_toc, pretty,
)

OUT = REPO_ROOT / "research_engine" / "output"
CHARTS = OUT / "charts"


def df_table(doc, df, cols, headers, widths=None, font=8, max_rows=40):
    rows = []
    for _, r in df.head(max_rows).iterrows():
        row = []
        for c in cols:
            v = r.get(c)
            if isinstance(v, float):
                v = round(v, 3)
            row.append(v)
        rows.append(row)
    add_table(doc, headers, rows, widths=widths, font_size=font)


def main() -> None:
    ladder = pd.read_csv(OUT / "engine_ladder.csv")
    traps = pd.read_csv(OUT / "trap_analysis.csv")
    trap_val = pd.read_csv(OUT / "trap_score_validation.csv")
    events = pd.read_csv(OUT / "event_impact.csv")
    proxies = pd.read_csv(OUT / "institutional_proxy.csv")
    dec = pd.read_csv(OUT / "decision_validation.csv")
    scored = pd.read_parquet(OUT / "trades_scored.parquet")
    n_total = len(scored)

    doc = Document()
    for sec in doc.sections:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(0.7)

    t = add_para(doc, "Market Lens Overall Strategy Engine", bold=True, size=24)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = add_para(doc, "Research layer 2 — combining zones, context, events, institutional proxies, "
                      "trap detection and trade planning into one evidence-tested decision engine",
                 italic=True, size=12)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = add_para(doc, f"Prepared 13 Aug 2026 · builds on the {n_total:,}-trade F&O backtest · 208 F&O stocks "
                      "+ NIFTY 50 + BANK NIFTY · 5 timeframes", size=10)
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w = add_para(doc, "For education, research and Market Lens scanner design only. Not investment advice. "
                      "Historical performance does not guarantee future results. Every weight in this engine is "
                      "calibrated on one year of data (in-sample); treat all results as upper bounds.",
                 italic=True, size=9)
    w.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    add_heading(doc, "Contents", 1)
    add_toc(doc)
    doc.add_page_break()

    # 1 executive summary
    add_heading(doc, "1. Executive Summary", 1)
    add_bullets(doc, [
        "VERDICT: the Overall Strategy Engine is PARTIALLY READY. Its core — setup quality x zone location x "
        "three validated trap gates x risk/reward — turned an aggregate −0.158R stream of 193,515 raw signals "
        "into +0.087R over 12,266 selected trades (profit factor 1.15, total +1,072R, max drawdown 143R vs "
        "31,040R unfiltered). The news layer is modest but usable; the institutional layer is NOT ready.",
        "Decisions separate cleanly: TAKE +0.088R, WAIT +0.136R, WATCH +0.091R versus NO TRADE −0.196R and "
        "AVOID −0.120R. The engine's job — isolating the tradeable minority — is achieved by GATES, not by a "
        "blended 0–100 score (a blended score tested non-monotonic and was demoted to a ranking aid).",
        "Zone location does most of the work: adding the direction-aligned-zone requirement alone moved "
        "recommended setups from −0.158R to breakeven while filtering 86% of signals.",
        "The trap study REFUTED most classic trap lore this year: overextension, sharp pre-signal moves, "
        "news-shock bars and 'shorting a bullish market' all REDUCED stop-outs within the same setups — a "
        "momentum year rewarded chasing. Only three traps validated: breakout without volume (+2.9pp "
        "stop-outs), longs into an overextended NIFTY (+1.7pp), and entries within 1 day of results (+1.6pp).",
        "Events: fighting a fresh result reaction was the real news trap (−0.055R stratified); post-result "
        "alignment and positive surprises gave mild tailwinds (+0.03 to +0.10R). Real news feeds remain "
        "unavailable historically — proxies were used and are labelled as such.",
        "Institutional proxies: OBV/AD alignment added nothing (−0.01R) and gating on them HURT the ladder; "
        "VWAP-side alignment intraday was the one strong proxy (+0.099R stratified). True institutional data "
        "(delivery %, OI build-up, FII/DII) was not available historically and is the top data gap before the "
        "institutional score can go live.",
    ])

    # 2 baseline
    add_heading(doc, "2. Previous Research Baseline", 1)
    add_para(doc, "Layer 1 (see FnO_Pattern_Research_Report.docx) established: most standalone setups lose after "
                  "costs; Gap-Up Continuation (+0.37R), Fresh Zone Touch (+0.12R) and demand-side confirmation "
                  "were the strongest rules; SMA50 alignment improved zone touches; weekly > daily >> 15m; "
                  "bearish setups broadly failed in the tested year; and location/context mattered more than "
                  "pattern names. This layer keeps that trade model identical (next-bar entry, ATR-buffered "
                  "stops, 2R target, time stop, 0.1% costs, both-touched = loss) and adds the context, event, "
                  "proxy, trap and decision layers on top of the same 816,937 simulated trades.")

    # 3 objective, 4 data, 5 assumptions
    add_heading(doc, "3. New Research Objective", 1)
    add_para(doc, "Decide, for every signal, whether it is a high-probability trade candidate, needs confirmation, "
                  "is risky, is a probable trap, or should not be traded — using only information available at the "
                  "signal bar, and measuring every added filter against the base it filters.")
    add_heading(doc, "4. Data Sources Used", 1)
    add_bullets(doc, [
        "Prices/volume: Yahoo Finance unadjusted OHLCV (daily 2y, 60m 1y, 15m/75m ~60 days).",
        "Zones: Market Lens zone engine, survivorship filter disabled for research, walked point-in-time.",
        "Earnings: yfinance earnings calendar — 4,658 result dates with EPS estimate/actual/surprise for 204 of "
        "208 stocks; next-result dates treated as known in advance (exchange intimations make this realistic).",
        "Market/sector context: NIFTY 50, BANK NIFTY and 10 sector indices (regime, extension, 20-day returns).",
        "NOT available historically, therefore NOT used: news feeds, delivery volume, option OI, FII/DII flows, "
        "India VIX history, market breadth. Where the request named them, transparent price/volume proxies were "
        "substituted and labelled as proxies.",
    ])
    add_heading(doc, "5. Assumptions and Limitations", 1)
    add_bullets(doc, [
        "One year, one broadly rising market. The refuted traps (overextension, chasing) may become real traps "
        "in a different regime — the single biggest caveat on this study.",
        "Setup-quality tiers and all weights are calibrated in-sample. Out-of-sample validation on a second year "
        "is the main item under 'needs more research'.",
        "Intraday (75m/15m) covers ~60 trading days. Index instruments carry no volume on Yahoo.",
        "No portfolio construction: trades are independent; drawdown figures assume 1R risk per trade, "
        "sequenced by exit date.",
    ])

    # 6 news methodology, 7 proxy methodology, 8 trap methodology
    add_heading(doc, "6. News / Result / Event Methodology", 1)
    add_para(doc, "Real headlines are unavailable historically, so 'news' is proxied by what news does to a chart: "
                  "gaps >= 2%, volume >= 2.5x its 20-bar average, single-bar moves >= 2 ATR, plus the earnings "
                  "calendar (dates, EPS surprise, next-day price reaction). Each condition is measured with a "
                  "composition-safe STRATIFIED delta: the effect is computed within each (setup, timeframe) cell "
                  "and pooled — a pooled contrast is confounded by signal mix and was shown to flip signs.")
    add_heading(doc, "7. Institutional Proxy Methodology", 1)
    add_para(doc, "Institutional buying/selling cannot be observed in OHLCV data. The proxies tested: OBV 20-bar "
                  "trend, Accumulation/Distribution 20-bar trend, intraday VWAP side, volume expansion/contraction, "
                  "accumulation-at-demand and distribution-at-supply composites, and liquidity deciles. The "
                  "Institutional Support Score is therefore explicitly a PROXY score; the report never claims "
                  "confirmed institutional flow.")
    add_heading(doc, "8. Trap Probability Methodology", 1)
    add_para(doc, "14 candidate trap conditions were defined up front, each with an early-detection rule. A "
                  "condition earns score points only if, within (setup, timeframe) strata, it BOTH raises the "
                  "stop-hit rate by >=0.5pp AND does not improve expectancy. Points are proportional to the "
                  "measured stop-hit uplift, scaled so all validated traps firing together = 100. Buckets: "
                  "0–20 low, 21–40 moderate, 41–60 caution, 61–80 high, 81–100 avoid.")

    # 9 scoring model
    add_heading(doc, "9. Overall Strategy Scoring Model", 1)
    add_para(doc, "Seven sub-scores are computed per signal (all 0–100): Technical Setup, Market Context, Stock "
                  "Context, News/Event, Institutional Proxy, Trap Probability and Risk/Reward. The FINAL DECISION "
                  "is made by evidence-ordered gates — setup×timeframe cohort quality, then location (aligned live "
                  "zone / zone or gap setup), then the three validated traps, then RR to the opposing zone. The "
                  "blended confidence score is retained only to rank candidates within a decision bucket, because "
                  "a blended-score decision tested non-monotonic. Weights the data refuted are ZEROED and listed "
                  "in §14 — including SMA200 alignment (~0 effect this year) and all overextension penalties "
                  "except the NIFTY-extension trap T13.")

    # 10-11 backtest + results
    add_heading(doc, "10. Backtesting Methodology — the Ablation Ladder", 1)
    add_para(doc, "Each filter layer is added cumulatively to the same base signals; nothing is assumed to help. "
                  "'Recommended setups' = the 14 setups shortlisted in layer 1.")
    add_heading(doc, "11. Overall Strategy Results", 1)
    for lab, title in [("recommended_setups", "All timeframes, recommended setups"),
                       ("recommended_daily_weekly_60m", "Daily + weekly + 60m only")]:
        add_para(doc, title, bold=True)
        sub = ladder[ladder.ladder == lab]
        df_table(doc, sub,
                 ["stage", "n", "pct_filtered_out", "expectancy_r", "win_rate_2r", "stop_hit_pct",
                  "profit_factor", "total_r", "max_drawdown_r"],
                 ["Stage", "N", "% filtered", "Exp (R)", "2R win %", "Stop %", "PF", "Total R", "Max DD (R)"],
                 widths=[1.9, 0.7, 0.7, 0.6, 0.7, 0.6, 0.5, 0.7, 0.8])
    add_bullets(doc, [
        "Rung 2 (zone context) does most of the work: −0.158R → ~0.000R while filtering 86%.",
        "Rungs 3–4 (market/news) trim little and cost little — they matter through the trap gates instead.",
        "Rung 5 (OBV/AD institutional gate) HURT (−0.027R) and is excluded from the engine's decision path — "
        "measured, reported, not shipped.",
        "Rung 7 (full engine: cohort quality + location + traps + RR) reaches +0.087R on 12,266 trades, "
        "PF 1.15. Filtering 93.7% of signals is the point: almost everything a raw scanner shows should not "
        "be traded.",
    ])
    add_para(doc, "Decision-category validation (all 816,937 signals):", bold=True)
    df_table(doc, dec.rename(columns={dec.columns[0]: "decision"}),
             ["decision", "n", "expectancy", "win2r", "stop"],
             ["Decision", "N", "Expectancy (R)", "2R win rate", "Stop-hit rate"], font=9)

    # 12-13 best combos
    add_heading(doc, "12. Best Bullish Strategy Combinations", 1)
    sc = scored[scored.final_decision == "TAKE"]
    bl = sc[sc.bullish_or_bearish == "bullish"].groupby(["setup_name", "timeframe"]).agg(
        n=("r_multiple", "size"), exp=("r_multiple", "mean"), win2r=("hit_2r", "mean")).reset_index()
    bl = bl[bl.n >= 30].sort_values("exp", ascending=False).head(12)
    bl["setup_name"] = bl["setup_name"].map(pretty)
    bl["win2r"] = (bl["win2r"] * 100).round(1)
    df_table(doc, bl, ["setup_name", "timeframe", "n", "exp", "win2r"],
             ["Setup (TAKE only)", "TF", "N", "Exp (R)", "2R win %"], font=9)
    add_heading(doc, "13. Best Bearish Strategy Combinations", 1)
    br = sc[sc.bullish_or_bearish == "bearish"].groupby(["setup_name", "timeframe"]).agg(
        n=("r_multiple", "size"), exp=("r_multiple", "mean"), win2r=("hit_2r", "mean")).reset_index()
    br = br[br.n >= 30].sort_values("exp", ascending=False).head(10)
    br["setup_name"] = br["setup_name"].map(pretty)
    br["win2r"] = (br["win2r"] * 100).round(1)
    df_table(doc, br, ["setup_name", "timeframe", "n", "exp", "win2r"],
             ["Setup (TAKE only)", "TF", "N", "Exp (R)", "2R win %"], font=9)
    add_para(doc, "Short-side TAKEs remain thin — consistent with layer 1's finding that bearish setups "
                  "struggled in a rising year. The engine correctly starves them of TAKE decisions.", italic=True)

    # 14 trap analysis
    doc.add_page_break()
    add_heading(doc, "14. Trap Analysis", 1)
    add_para(doc, "Pooled vs stratified matters: several conditions look protective or dangerous in a pooled "
                  "contrast purely because of WHICH setups/timeframes they fire on. The stratified column is the "
                  "honest one.", italic=True)
    tshow = traps[["trap_id", "trap", "n", "stop_uplift_pp_pooled", "stop_uplift_pp_stratified",
                   "delta_expectancy_stratified", "false_breakout_pct", "worst_timeframe", "early_detection"]]
    df_table(doc, tshow,
             list(tshow.columns),
             ["ID", "Trap condition", "N", "Stop Δpp (pooled)", "Stop Δpp (strat.)", "Exp Δ (strat.)",
              "False brk %", "Worst TF", "Early detection"],
             widths=[0.35, 1.9, 0.55, 0.7, 0.7, 0.6, 0.6, 0.55, 1.7], font=7)
    add_bullets(doc, [
        "VALIDATED traps (score points): T4 breakout without volume (+2.9pp stops), T13 longs into NIFTY "
        "extended >2 ATR (+1.7pp, −0.072R), T9 signal within 1 day of results (+1.6pp).",
        "REFUTED this year (zero points, reported): T8 sharp pre-signal move (−9.7pp stops — momentum "
        "continued), T6 overextended entry (−6.4pp), T12 news-shock bar (−4.4pp), T7 short-in-bull-market "
        "(−3.6pp), T1/T2 trading into an opposing zone (−1.4/−0.8pp), T14 recently-broken zone, T5 counter-"
        "SMA200 (~0), T10 stale zones (~0), T3 over-tested level, T11 low liquidity (+0.3pp, below threshold).",
        "Score validation (composition-safe): score>0 → +2.1pp stop-outs / −0.030R; ≥40 → +2.9pp; "
        "≥60 → +6.2pp / −0.040R. Weak but monotonic — an honest, narrow trap layer.",
    ])
    df_table(doc, trap_val, list(trap_val.columns),
             ["Comparison", "N flagged", "Stratified stop Δpp", "Stratified exp Δ"], font=9)

    # 15 news impact
    add_heading(doc, "15. News / Result Impact Analysis", 1)
    ev = events[["condition", "n", "expectancy_r", "delta_r_stratified", "stop_hit_pct"]]
    df_table(doc, ev, list(ev.columns),
             ["Condition", "N", "Exp (R)", "Δ vs same setup/TF (strat.)", "Stop %"],
             widths=[2.9, 0.6, 0.6, 1.1, 0.6], font=8)
    add_bullets(doc, [
        "Usable: do not FIGHT a fresh result reaction (−0.055R); mild tailwind from trading WITH post-result "
        "alignment (+0.035 to +0.049R) and from positive surprises (+0.031R); post-result days were fine to "
        "trade (+0.099R stratified) — the risk sits BEFORE results, not after.",
        "Pre-result gate (T9) is small but real; the engine flags rather than forbids day-2 proximity.",
    ])

    # 16 institutional proxies
    add_heading(doc, "16. Institutional Proxy Analysis", 1)
    px = proxies[["condition", "n", "expectancy_r", "delta_r_stratified", "stop_hit_pct"]]
    df_table(doc, px, list(px.columns),
             ["Proxy", "N", "Exp (R)", "Δ vs same setup/TF (strat.)", "Stop %"],
             widths=[2.9, 0.6, 0.6, 1.1, 0.6], font=8)
    add_bullets(doc, [
        "VWAP-side alignment is the standout proxy (+0.099R, intraday) — recommend as an intraday confirmation "
        "column in Market Lens.",
        "OBV/AD trend alignment: ~zero to negative, and gating on them hurt the ladder — do not ship as a gate.",
        "Volume expansion (+0.048R) and pre-breakout contraction (+0.017R) remain the best volume signals; "
        "bottom-decile liquidity is a mild negative (−0.022R).",
        "The Institutional Support Score therefore rests on weak legs today. It becomes meaningful only with "
        "real data: NSE delivery %, F&O OI build-up classification, FII/DII flows. Until then it is displayed "
        "as 'proxy' and given no veto power.",
    ])

    # 17-18 trade planning
    add_heading(doc, "17. Risk/Reward and Trade Planning Rules", 1)
    add_bullets(doc, [
        "Stops: zone trades at distal ± 0.1×ATR; breakouts at level ± 1.0×ATR; candles at pattern extreme ± "
        "0.1×ATR (as backtested).",
        "Targets: T1 = 1R (partial book / move stop to entry), T2 = 2R primary (backtested exit), T3 = 3R "
        "stretch; where a live opposing zone sits closer than 1.5R, the trade is WAIT/NO TRADE.",
        "RR gate evidence: rr_to_opposing >= 1.5 (or no obstacle) is required for TAKE.",
        "Position sizing: fixed rupee risk per trade (the R framework); REDUCE SIZE bucket = half risk. "
        "Risk 1–4% of entry preferred; >6% skipped.",
        "Trailing (not backtested — forward proposal): after 1R, stop to entry; after 2R, trail under the most "
        "recent completed zone or 20-EMA, whichever is nearer.",
    ])
    add_heading(doc, "18. Entry / Stop-Loss / Target Logic (per entry type)", 1)
    add_table(doc, ["Entry type", "Entry", "Stop", "Notes"], [
        ["Fresh zone touch (Type-1)", "Limit at proximal", "Distal − 0.1 ATR", "No favourable credit on fill bar (conservative fills)"],
        ["Zone confirmation bounce", "Next open after confirmation close", "Distal − 0.1 ATR", "Demand side validated; supply side weak this year"],
        ["Gap continuation", "Next open (or same-open live)", "Prior bar low − 0.1 ATR", "Works unfiltered; zone context optional bonus"],
        ["Pattern breakout", "Next open after confirmed close", "Level ∓ 1.0 ATR", "Volume expansion REQUIRED (T4)"],
        ["Pullback (EMA20/SMA50)", "Next open after reclaim close", "Bounce bar extreme ∓ 0.1 ATR", "75m rejection-at-supply variant for shorts"],
    ], widths=[1.5, 1.7, 1.4, 2.5], font_size=8)

    # 19-20 stock/sector
    add_heading(doc, "19. Stock-Specific Findings (engine TAKE trades)", 1)
    st = sc.groupby("symbol").agg(n=("r_multiple", "size"), exp=("r_multiple", "mean")).reset_index()
    st = st[st.n >= 25]
    top = st.sort_values("exp", ascending=False).head(8)
    bot = st.sort_values("exp").head(5)
    df_table(doc, pd.concat([top, bot]), ["symbol", "n", "exp"],
             ["Symbol (top 8 / bottom 5, n>=25)", "TAKE trades", "Expectancy (R)"], font=9)
    add_heading(doc, "20. Sector-Specific Findings (engine TAKE trades)", 1)
    se = sc.groupby("sector").agg(n=("r_multiple", "size"), exp=("r_multiple", "mean"),
                                  win2r=("hit_2r", "mean")).reset_index().sort_values("exp", ascending=False)
    se["win2r"] = (se["win2r"] * 100).round(1)
    df_table(doc, se, ["sector", "n", "exp", "win2r"], ["Sector", "N", "Exp (R)", "2R win %"], font=9)

    # 21-22 examples
    doc.add_page_break()
    add_heading(doc, "21. Example Winning Trades (engine TAKE)", 1)
    add_para(doc, "Green = entry, red dashed = stop, purple dotted = 2R target; grey vertical = signal bar.",
             italic=True, size=9)
    for png in sorted(CHARTS.glob("engine_take_win_*.png")):
        doc.add_picture(str(png), width=Inches(6.7))
        add_para(doc, png.stem.replace("_", " "), size=9, italic=True)
    add_heading(doc, "22. Example Losing and Trap Trades", 1)
    for png in sorted(CHARTS.glob("engine_take_loss_*.png")) + sorted(CHARTS.glob("engine_avoid_trap_*.png")):
        doc.add_picture(str(png), width=Inches(6.7))
        add_para(doc, png.stem.replace("_", " "), size=9, italic=True)
    add_para(doc, "TAKE losses are included deliberately: even the filtered stream stops out ~54% of the time — "
                  "the engine's edge is expectancy over many trades, never certainty on one.", italic=True)

    # 23 implementation
    doc.add_page_break()
    add_heading(doc, "23. Market Lens Implementation Recommendations", 1)
    add_table(doc, ["Feature", "Evidence", "Build", "Complexity", "Notes"], [
        ["Decision gates (cohort tier × zone location × traps × RR) as a 'Strategy Engine' column set",
         "+0.087R vs −0.158R base; monotone decisions", "NOW", "Medium",
         "Reuses existing zone engine + new gap/trend scanners from layer 1"],
        ["Trap flags T4/T13/T9 with per-flag tooltips",
         "+1.6 to +2.9pp stop-outs each (stratified)", "NOW", "Low",
         "Volume check, NIFTY extension, earnings-calendar join (calendar already in app)"],
        ["RR-to-opposing-zone column + 1.5R gate",
         "Gate condition of the TAKE bucket", "NOW", "Low",
         "Distances already computed by the zone engine"],
        ["Trap Probability Score 0–100 UI (3 validated traps only)",
         "Monotone but weak (+2–6pp)", "NOW (labelled 'narrow')", "Low",
         "Show flags, not just the number"],
        ["VWAP-side confirmation column (intraday)",
         "+0.099R stratified", "NOW", "Low", "Intraday scans only"],
        ["News/Event score", "Small effects; avoid-fighting-reaction rule",
         "LATER (partial: result-proximity flag NOW)", "Medium", "Needs a real news feed for more"],
        ["Institutional Support Score", "OBV/AD refuted; VWAP only",
         "LATER — needs delivery %, OI build-up, FII/DII pipelines", "High",
         "Display as PROXY; no veto power until real data"],
        ["Out-of-sample re-validation harness", "All weights are in-sample",
         "BEFORE going live", "Medium", "Re-run research_engine/harness on a second year / rolling window"],
    ], widths=[2.0, 1.6, 1.1, 0.8, 1.6], font_size=8)

    # 24-26 verdict
    add_heading(doc, "24. Final Decision — Is the Engine Ready?", 1)
    add_para(doc, "PARTIALLY READY.", bold=True)
    add_bullets(doc, [
        "READY NOW: setup×timeframe quality tiers; zone-location gate; traps T4/T13/T9; RR-to-opposing-zone "
        "gate; decision categories TAKE / WAIT / WATCH / REDUCE SIZE / AVOID / NO TRADE; VWAP intraday "
        "confirmation; entry/stop/target computation as specified in §18.",
        "NEEDS MORE RESEARCH: out-of-sample validation (top priority — every weight is in-sample); regime "
        "dependence of the refuted traps; supply-side (short) rules; the blended confidence score as anything "
        "more than a sort key; 75m/15m rules on a full year of data.",
        "DO NOT IMPLEMENT: OBV/AD gating; market-direction filters on bullish signals (tested harmful); "
        "generic 'trap lore' penalties the data refuted; any claim of confirmed institutional activity.",
        "DATA REQUIRED BEFORE FULL LIVE: NSE delivery %, F&O OI with build-up classification, FII/DII flows, "
        "a real news feed with timestamps, India VIX history, market breadth.",
    ])
    add_heading(doc, "25. What Still Needs More Research", 1)
    add_bullets(doc, [
        "A second (ideally bearish/sideways) year to re-test every weight — especially the refuted traps.",
        "Walk-forward calibration: recompute cohort tiers quarterly instead of a single full-year fit.",
        "Exit engineering: trailing rules, partial booking at 1R, and opposing-zone targets were specified but "
        "not backtested as exits.",
        "Options overlay: translating TAKE candidates into strikes/expiries/IV filters (out of scope here).",
    ])
    add_heading(doc, "26. Final Checklist of Scanner Rules", 1)
    add_table(doc, ["#", "Rule", "Decision"], [
        ["E1", "Engine TAKE = qualified setup + aligned live zone/gap + trap<40 + RR>=1.5", "Build now"],
        ["E2", "Trap flags T4 / T13 / T9 shown as chips with tooltips", "Build now"],
        ["E3", "RR-to-opposing-zone column + gate", "Build now"],
        ["E4", "VWAP-side intraday confirmation column", "Build now"],
        ["E5", "Result-proximity flag (<=2 days) from the earnings calendar", "Build now"],
        ["E6", "News/Event score beyond result proximity", "Later (needs news feed)"],
        ["E7", "Institutional Support Score with real delivery/OI/FII-DII data", "Later (needs data)"],
        ["E8", "OBV/AD alignment gates", "Do not build (tested harmful)"],
        ["E9", "Generic overextension/'don't short a bull market' filters", "Do not build (refuted this year)"],
    ], widths=[0.4, 4.9, 1.6], font_size=9)

    dest = OUT / "Overall_Strategy_Engine_Report.docx"
    doc.save(dest)
    print("saved", dest)


if __name__ == "__main__":
    main()
